#!/usr/bin/env python3
"""
Gera 2304335-relatorio-intercalar-lei.docx a partir do template UAb e do
conteúdo de docs/report/relatorio-intercalar.md.

Uso: py -3 scripts/build_docx.py
"""

import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Paths ────────────────────────────────────────────────────────────────────

TEMPLATE = Path(
    r"C:\Users\jrzab\iCloudDrive\iCloud~md~obsidian\LifeOS"
    r"\02 - Study\21184-Projeto-Engenharia-Informatica"
    r"\attachments\Relatório LEI - Modelo.docx"
)
MARKDOWN = Path(
    r"C:\Users\jrzab\IdeaProjects\21184-ProjetoLEI"
    r"\docs\report\relatorio-intercalar.md"
)
OUTPUT = Path(
    r"C:\Users\jrzab\IdeaProjects\21184-ProjetoLEI"
    r"\docs\report\2304335-relatorio-intercalar-lei.docx"
)

# ─── Cover substitutions ──────────────────────────────────────────────────────

COVER_SUBS = {
    '<Título do Relatório>':       'Musical Theory Trainer',
    '<Subtítulo do Relatório>':    'Relatório Intercalar',
    'ReLATÓRIO FINAL':             'RELATÓRIO INTERCALAR',
    '<Nome e número do aluno>':    'Daniel Junior - 2304335',
    '<Nome do orientador>':        'Pedro Duarte Pestana',
    '<Data>':                      'Maio 2026',
    '<Inserir índice de figuras>': '[Índice de figuras - gerar no Word: References → Insert Table of Figures]',
    '<Inserir índice de tabelas>': '[Índice de tabelas - gerar no Word: References → Insert Table of Figures]',
}

RESUMO_PARAGRAPHS = [
    ('O Musical Theory Trainer é uma aplicação web de treino auditivo para teoria '
     'musical. O utilizador ouve intervalos, escalas e acordes e responde tocando '
     'as notas num teclado virtual ou num controlador MIDI físico. A aplicação '
     'avalia automaticamente, regista o progresso e adapta a dificuldade ao '
     'desempenho do utilizador ao longo das sessões.'),
    ('Este relatório intercalar documenta o desenho completo do sistema - '
     'arquitectura em 3 camadas, modelo de dados, três algoritmos principais '
     '(dificuldade adaptativa, identificação de padrões fracos, geração procedural), '
     'wireframes e 19 ADRs - e o estado de implementação a 2 de maio de 2026: '
     'backend completo com 230 testes a passar, frontend especificado e agendado '
     'para as semanas 9-12.'),
]

PLACEHOLDER_CHAPTERS = [
    ('Capítulo 4 - Testes',
     'A desenvolver no relatório final (entrega: 24 de junho de 2026). '
     'Incluirá: exemplos de funcionamento normal com capturas de ecrã, '
     'testes unitários e de integração, testes de desempenho e casos-limite.'),
    ('Capítulo 5 - Conclusões',
     'A desenvolver no relatório final (entrega: 24 de junho de 2026). '
     'Incluirá: reflexão sobre resultados face aos objectivos, dificuldades '
     'e limitações, melhorias possíveis e trabalho futuro.'),
]

BIBLIOGRAPHY = [
    'Alur, D., Crupi, J., & Malks, D. (2003). Core J2EE Patterns: Best Practices and Design Strategies (2nd ed.). Prentice Hall.',
    'Anthropic. (2025). Claude Code [Ferramenta de assistência com IA]. Utilizada ao longo do desenvolvimento para pesquisa de alternativas técnicas, sugestões de boas práticas, exemplos de código e revisão de documentação. https://claude.ai/code',
    'Csikszentmihalyi, M. (1990). Flow: The psychology of optimal experience. Harper & Row.',
    'Evans, E. (2003). Domain-Driven Design: Tackling Complexity in the Heart of Software. Addison-Wesley.',
    'Fowler, M. (2002). Patterns of Enterprise Application Architecture. Addison-Wesley.',
    'Gamma, E., Helm, R., Johnson, R., & Vlissides, J. (1994). Design Patterns: Elements of Reusable Object-Oriented Software. Addison-Wesley.',
    'Internet Engineering Task Force. (2016). RFC 7807 - Problem Details for HTTP APIs. https://www.rfc-editor.org/rfc/rfc7807',
    'Spring Team. (2024). Spring Boot 3.3.0 Reference Documentation. https://docs.spring.io/spring-boot/docs/3.3.0/reference/html/',
    'Vygotsky, L. S. (1978). Mind in society: The development of higher psychological processes. Harvard University Press.',
    'World Wide Web Consortium. (2021). Web Audio API - W3C Recommendation. https://www.w3.org/TR/webaudio/',
    "World Wide Web Consortium. (2015). Web MIDI API - W3C Editor's Draft. https://www.w3.org/TR/webmidi/",
]

# ─── Inline formatting ────────────────────────────────────────────────────────

def parse_runs(text):
    """
    Returns list of (text, bold, italic, code).
    Handles **bold**, *italic*, _italic_, `code`, and strips [links](url).
    """
    text = re.sub(r'!\[[^\]]*\]\([^)]+\)', '', text)     # strip images
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)  # links → text

    parts = []
    pattern = re.compile(r'\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|_(.+?)_|\*(.+?)\*|`(.+?)`')
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            parts.append((text[pos:m.start()], False, False, False))
        if m.group(1):
            parts.append((m.group(1), True, True, False))
        elif m.group(2):
            parts.append((m.group(2), True, False, False))
        elif m.group(3):
            parts.append((m.group(3), False, True, False))
        elif m.group(4):
            parts.append((m.group(4), False, True, False))
        elif m.group(5):
            parts.append((m.group(5), False, False, True))
        pos = m.end()
    if pos < len(text):
        parts.append((text[pos:], False, False, False))
    return [(t, b, i, c) for t, b, i, c in parts if t.strip() != '' or not (b or i or c)]


def apply_runs(paragraph, text):
    """Add runs with inline formatting to a paragraph."""
    for (t, bold, italic, code) in parse_runs(text):
        if not t:
            continue
        run = paragraph.add_run(t)
        run.bold = bold
        run.italic = italic
        if code:
            run.font.name = 'Courier New'
            run.font.size = Pt(9)


# ─── Paragraph / table helpers ────────────────────────────────────────────────

def add_para(doc, text, style):
    p = doc.add_paragraph(style=style)
    apply_runs(p, text)
    return p


def insert_para_after(ref_para, doc, text, style):
    """Insert a paragraph immediately after ref_para."""
    new_p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), style)
    pPr.append(pStyle)
    new_p.append(pPr)
    r_el = OxmlElement('w:r')
    t_el = OxmlElement('w:t')
    t_el.text = text
    t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r_el.append(t_el)
    new_p.append(r_el)
    ref_para._element.addnext(new_p)


def add_table_md(doc, rows):
    """Add a Word table from markdown rows (first row = header)."""
    if not rows or not rows[0]:
        return
    ncols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=ncols)
    tbl.style = 'Normal Table'
    for r_i, row in enumerate(rows):
        for c_i in range(ncols):
            cell_text = row[c_i].strip() if c_i < len(row) else ''
            cell = tbl.cell(r_i, c_i)
            cell.text = ''
            p = cell.paragraphs[0]
            runs = parse_runs(cell_text)
            for (t, bold, italic, code) in runs:
                if not t:
                    continue
                run = p.add_run(t)
                run.bold = True if r_i == 0 else bold
                run.italic = italic
                if code:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(8)
    doc.add_paragraph()


# ─── Cover page ───────────────────────────────────────────────────────────────

def fix_cover(doc):
    resumo_para = None
    for para in doc.paragraphs:
        full = ''.join(r.text for r in para.runs)
        for placeholder, replacement in COVER_SUBS.items():
            if placeholder in full:
                for run in para.runs:
                    run.text = ''
                if para.runs:
                    para.runs[0].text = replacement
                break
        if '<Descreva de forma resumidamente o projeto>' in full:
            resumo_para = para

    if resumo_para:
        # Replace placeholder with resumo text, change to Normal style
        resumo_para.style = doc.styles['Normal']
        for run in resumo_para.runs:
            run.text = ''
        if resumo_para.runs:
            resumo_para.runs[0].text = RESUMO_PARAGRAPHS[0]
        else:
            resumo_para.add_run(RESUMO_PARAGRAPHS[0])
        # Insert second resumo paragraph after the first
        insert_para_after(resumo_para, doc, RESUMO_PARAGRAPHS[1], 'Normal')
        print("  Resumo filled.")
    else:
        print("  WARNING: Resumo placeholder not found.")


# ─── Remove template chapter content ─────────────────────────────────────────

def remove_template_chapters(doc):
    """Remove all elements from the first Ttulo1 heading onwards."""
    first_h1 = None
    for para in doc.paragraphs:
        if para.style.style_id == 'Ttulo1' and para.text.strip():
            first_h1 = para._element
            break

    if first_h1 is None:
        print("  WARNING: No Ttulo1 found to remove.")
        return

    body = first_h1.getparent()
    removing = False
    to_remove = []
    for child in list(body):
        if child is first_h1:
            removing = True
        if removing:
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if tag != 'sectPr':
                to_remove.append(child)

    for el in to_remove:
        body.remove(el)
    print(f"  Removed {len(to_remove)} template body elements.")


# ─── Markdown parser ──────────────────────────────────────────────────────────

def iter_md_blocks(md_text):
    """
    Yield (kind, data) tuples from markdown.
    Kinds: heading1, heading2, heading3, para, bullet, code, table
    Skips: header block (before first ## Capítulo), section 2.9, ---, images.
    """
    lines = md_text.splitlines()
    started = False
    in_code = False
    code_buf = []
    in_table = False
    table_rows = []
    skip_section = False   # skip 2.9 references

    i = 0
    while i < len(lines):
        line = lines[i]

        # Code fence toggle
        stripped = line.strip()
        if stripped.startswith('```'):
            if in_code:
                in_code = False
                yield ('code', '\n'.join(code_buf))
                code_buf = []
            else:
                if in_table:
                    yield ('table', table_rows)
                    table_rows = []
                    in_table = False
                in_code = True
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # Table lines
        if stripped.startswith('|'):
            if re.match(r'^\s*\|[\s\-:|]+\|\s*$', line):
                i += 1
                continue
            in_table = True
            cells = [c.strip() for c in stripped.strip('|').split('|')]
            table_rows.append(cells)
            i += 1
            continue
        else:
            if in_table:
                yield ('table', table_rows)
                table_rows = []
                in_table = False

        # Wait for first ## Capítulo
        if not started:
            if re.match(r'^## Capítulo', line):
                started = True
            else:
                i += 1
                continue

        # Headings reset skip_section
        if re.match(r'^#{1,4} ', line):
            skip_section = False

        if skip_section:
            i += 1
            continue

        # Headings
        if re.match(r'^## Capítulo', line):
            yield ('heading1', line[3:].strip())
            i += 1
            continue
        if re.match(r'^## ', line):
            yield ('heading1', line[3:].strip())
            i += 1
            continue
        if re.match(r'^### ', line):
            text = line[4:].strip()
            if re.match(r'^2\.9\b', text):
                skip_section = True
                i += 1
                continue
            yield ('heading2', text)
            i += 1
            continue
        if re.match(r'^#### ', line):
            yield ('heading3', line[5:].strip())
            i += 1
            continue

        # Skip separators and empty
        if stripped in ('---', ''):
            i += 1
            continue

        # Image lines
        if re.match(r'^!\[', stripped):
            m_img = re.match(r'^!\[([^\]]*)\]\(([^)]+)\)', stripped)
            if m_img:
                yield ('image', (m_img.group(2), m_img.group(1)))
            i += 1
            continue

        # Indented pseudocode (4 spaces, not a list)
        if re.match(r'^    \S', line) and not re.match(r'^    -', line):
            code_lines = []
            while i < len(lines) and (re.match(r'^    ', lines[i]) or lines[i].strip() == ''):
                if lines[i].strip():
                    code_lines.append(lines[i][4:])
                else:
                    if code_lines:
                        code_lines.append('')
                i += 1
            yield ('code', '\n'.join(code_lines).rstrip())
            continue

        # Bullets (with optional indentation)
        m = re.match(r'^(\s*)- (.*)', line)
        if m:
            yield ('bullet', m.group(2).strip())
            i += 1
            continue

        # Regular paragraph
        if stripped:
            yield ('para', stripped)
        i += 1

    # Flush
    if in_table and table_rows:
        yield ('table', table_rows)
    if in_code and code_buf:
        yield ('code', '\n'.join(code_buf))


# ─── Insert content ───────────────────────────────────────────────────────────

def insert_content(doc, md_text, md_dir):
    count = {'heading1': 0, 'heading2': 0, 'heading3': 0,
             'para': 0, 'bullet': 0, 'code': 0, 'table': 0, 'image': 0}
    for (kind, data) in iter_md_blocks(md_text):
        count[kind] = count.get(kind, 0) + 1
        if kind == 'heading1':
            add_para(doc, data, 'Ttulo1')
        elif kind == 'heading2':
            add_para(doc, data, 'Ttulo2')
        elif kind == 'heading3':
            add_para(doc, data, 'Ttulo3')
        elif kind == 'para':
            add_para(doc, data, 'Normal')
        elif kind == 'bullet':
            add_para(doc, data, 'Listacommarcas')
        elif kind == 'code':
            p = doc.add_paragraph(style='HTMLpr-formatado')
            run = p.add_run(data)
            run.font.name = 'Courier New'
            run.font.size = Pt(8)
        elif kind == 'table':
            add_table_md(doc, data)
        elif kind == 'image':
            rel_path, alt = data
            img_path = (md_dir / rel_path).resolve()
            if img_path.exists():
                p = doc.add_paragraph()
                p.alignment = 1  # CENTER
                run = p.add_run()
                run.add_picture(str(img_path), width=Cm(14))
                # Caption
                cap = doc.add_paragraph(style='Legenda')
                cap.add_run(alt)
            else:
                add_para(doc, f'[Imagem não encontrada: {rel_path}]', 'Normal')
    print(f"  Inserted: {count}")


def insert_placeholders(doc):
    for (title, text) in PLACEHOLDER_CHAPTERS:
        add_para(doc, title, 'Ttulo1')
        add_para(doc, text, 'Normal')
    print("  Placeholder chapters 4-5 inserted.")


def insert_bibliography(doc):
    add_para(doc, 'Bibliografia', 'Ttulo1')
    for entry in BIBLIOGRAPHY:
        add_para(doc, entry, 'Normal')
    print(f"  Bibliography: {len(BIBLIOGRAPHY)} entries.")


# ─── Main ─────────────────────────────────────────────────────────────────────

def main():
    if not TEMPLATE.exists():
        print(f"ERROR: template not found: {TEMPLATE}")
        sys.exit(1)
    if not MARKDOWN.exists():
        print(f"ERROR: markdown not found: {MARKDOWN}")
        sys.exit(1)

    print(f"Loading template...")
    doc = Document(str(TEMPLATE))

    print("Fixing cover page...")
    fix_cover(doc)

    print("Removing template chapter content...")
    remove_template_chapters(doc)

    print("Parsing and inserting report content...")
    md_text = MARKDOWN.read_text(encoding='utf-8')
    insert_content(doc, md_text, MARKDOWN.parent)

    print("Inserting placeholder chapters 4 and 5...")
    insert_placeholders(doc)

    print("Inserting bibliography...")
    insert_bibliography(doc)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    print(f"\nDone. Output: {OUTPUT}")
    print("Open in Word, check formatting, generate PDF via File -> Export.")


if __name__ == '__main__':
    main()
